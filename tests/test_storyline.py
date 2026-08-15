import sys
import unittest
from unittest.mock import patch

from analytics_platform.storyline import (
    DocxRendererUnavailable, StorylineContent, StorylineTurn, CodeAppendixEntry,
    assemble_storyline, render_markdown, render_docx,
)


def _msg(answer_id, question="Q", answer="A", facts=None, caveats=None,
         queries_run=None, python_cells=None, produced_df_label=""):
    return {
        "answer_id": answer_id, "question": question, "answer": answer,
        "facts": facts or [], "caveats": caveats or [],
        "queries_run": queries_run or [], "python_cells": python_cells or [],
        "produced_df_label": produced_df_label, "created_at": "2026-08-15T00:00:00Z",
    }


class TestAssembleStoryline(unittest.TestCase):
    def test_only_selected_turns_appear_in_order(self):
        conv = {"id": "c1", "title": "Test", "messages": [
            _msg("a1", question="First"), _msg("a2", question="Second"),
            _msg("a3", question="Third"),
        ]}
        content = assemble_storyline(conv, ["a3", "a1"])
        self.assertEqual([t.question for t in content.turns], ["First", "Third"])

    def test_empty_selection_yields_empty_content(self):
        conv = {"id": "c1", "title": "Test", "messages": [_msg("a1")]}
        content = assemble_storyline(conv, [])
        self.assertEqual(content.turns, [])
        self.assertEqual(content.code_appendix, [])
        self.assertEqual(content.estimated_tokens, 0)
        self.assertFalse(content.over_budget)

    def test_selected_sql_turn_adds_its_query_as_non_dependency_appendix_entry(self):
        conv = {"id": "c1", "title": "Test", "messages": [
            _msg("a1", queries_run=["SELECT 1"], produced_df_label="df_1"),
        ]}
        content = assemble_storyline(conv, ["a1"])
        self.assertEqual(len(content.code_appendix), 1)
        entry = content.code_appendix[0]
        self.assertEqual(entry.kind, "sql")
        self.assertEqual(entry.code, "SELECT 1")
        self.assertEqual(entry.source_answer_id, "a1")
        self.assertFalse(entry.is_dependency)

    def test_selected_python_turn_pulls_in_unselected_producing_turn_as_dependency(self):
        conv = {"id": "c1", "title": "Test", "messages": [
            _msg("a1", question="SQL turn", queries_run=["SELECT revenue FROM events"],
                 produced_df_label="df_1"),
            _msg("a2", question="Python turn",
                 python_cells=[{"code": "result = df_1['revenue'].sum()",
                                "df_label": "df_1", "result_summary": 6}]),
        ]}
        content = assemble_storyline(conv, ["a2"])  # a1 NOT selected
        self.assertEqual([t.question for t in content.turns], ["Python turn"])
        kinds = {(e.kind, e.source_answer_id, e.is_dependency) for e in content.code_appendix}
        self.assertIn(("python", "a2", False), kinds)
        self.assertIn(("sql", "a1", True), kinds)

    def test_dependency_turn_that_is_also_selected_is_not_double_marked(self):
        conv = {"id": "c1", "title": "Test", "messages": [
            _msg("a1", queries_run=["SELECT 1"], produced_df_label="df_1"),
            _msg("a2", python_cells=[{"code": "result = 1", "df_label": "df_1",
                                      "result_summary": 1}]),
        ]}
        content = assemble_storyline(conv, ["a1", "a2"])
        sql_entries = [e for e in content.code_appendix if e.kind == "sql"]
        self.assertEqual(len(sql_entries), 1)
        self.assertFalse(sql_entries[0].is_dependency)

    def test_recycled_df_label_resolves_to_the_nearest_preceding_producer(self):
        """C1 regression. df_label is a recyclable display label, not an identity:
        the in-memory frame cache is LRU-bounded and reset on API restart, so one
        persisted conversation can contain two turns that both produced 'df_1'. A
        last-wins {label: message} map attributed the Python turn to whichever turn
        reused the label LAST -- here a query that ran *after* it, over unrelated
        data. A Python cell can only read a frame an earlier turn cached, so the
        producer must be the nearest one *preceding* the Python turn."""
        conv = {"id": "c1", "title": "Test", "messages": [
            _msg("ans_A", question="EU revenue",
                 queries_run=["SELECT revenue FROM events WHERE region='EU'"],
                 produced_df_label="df_1"),
            _msg("ans_B", question="Python turn",
                 python_cells=[{"code": "result = df_1['revenue'].sum()",
                                "df_label": "df_1", "result_summary": 6}]),
            _msg("ans_H", question="refunds", queries_run=["SELECT refunds FROM tickets"],
                 produced_df_label="df_1"),
        ]}
        content = assemble_storyline(conv, ["ans_B"])  # only the Python turn
        deps = [e for e in content.code_appendix if e.is_dependency]
        self.assertEqual(len(deps), 1)
        self.assertEqual(deps[0].source_answer_id, "ans_A")
        self.assertEqual(deps[0].code, "SELECT revenue FROM events WHERE region='EU'")
        self.assertEqual(content.unresolved_dependency_count, 0)
        codes = [e.code for e in content.code_appendix]
        self.assertNotIn("SELECT refunds FROM tickets", codes)

    def test_unresolvable_df_label_yields_a_visible_note_not_silence(self):
        """I3. produced_df_label was added by ALTER TABLE, so every pre-existing row
        backfills to "" -- for those conversations no Python turn's SQL dependency is
        resolvable. Dropping it silently would ship a Code Appendix with Python over
        an undefined df_1 and no hint that anything is missing."""
        conv = {"id": "c1", "title": "Test", "messages": [
            _msg("a1", question="Python turn",
                 python_cells=[{"code": "result = df_1['revenue'].sum()",
                                "df_label": "df_1", "result_summary": 6}]),
        ]}
        content = assemble_storyline(conv, ["a1"])
        notes = [e for e in content.code_appendix if e.kind == "note"]
        self.assertEqual(len(notes), 1)
        self.assertEqual(notes[0].label, "df_1")
        self.assertTrue(notes[0].is_dependency)
        self.assertIn("not recorded", notes[0].code)
        self.assertEqual(content.unresolved_dependency_count, 1)

    def test_unresolvable_note_is_emitted_once_per_label(self):
        conv = {"id": "c1", "title": "Test", "messages": [
            _msg("a1", python_cells=[{"code": "x = 1", "df_label": "df_9",
                                      "result_summary": 1}]),
            _msg("a2", python_cells=[{"code": "y = 2", "df_label": "df_9",
                                      "result_summary": 2}]),
        ]}
        content = assemble_storyline(conv, ["a1", "a2"])
        self.assertEqual(len([e for e in content.code_appendix if e.kind == "note"]), 1)
        self.assertEqual(content.unresolved_dependency_count, 1)

    def test_token_estimate_and_budget_flag(self):
        conv = {"id": "c1", "title": "Test", "messages": [
            _msg("a1", question="Q" * 100, answer="A" * 100),
        ]}
        content = assemble_storyline(conv, ["a1"])
        self.assertGreater(content.estimated_tokens, 0)
        self.assertFalse(content.over_budget)

        big_conv = {"id": "c1", "title": "Test", "messages": [
            _msg("a1", answer="x" * 250_000),
        ]}
        big_content = assemble_storyline(big_conv, ["a1"])
        self.assertTrue(big_content.over_budget)

    def test_render_markdown_includes_turns_and_dependency_annotated_appendix(self):
        content = StorylineContent(
            conversation_title="Q3 Funnel Review",
            turns=[StorylineTurn(answer_id="a1", question="Why did signups drop?",
                                  answer="Signups dropped 12% after the consent page.",
                                  facts=["computed via SQL"], caveats=[],
                                  created_at="2026-08-15T00:00:00Z")],
            code_appendix=[
                CodeAppendixEntry(label="df_1", kind="sql", code="SELECT 1",
                                  source_answer_id="a0", is_dependency=True),
                CodeAppendixEntry(label="df_1", kind="python", code="result = 1",
                                  source_answer_id="a1", is_dependency=False),
            ],
            estimated_tokens=42, over_budget=False,
        )
        md = render_markdown(content)
        self.assertIn("# Q3 Funnel Review", md)
        self.assertIn("Why did signups drop?", md)
        self.assertIn("Signups dropped 12%", md)
        self.assertIn("```sql", md)
        self.assertIn("```python", md)
        self.assertIn("(included as a dependency of df_1)", md)

    def test_render_docx_produces_a_valid_document_with_turns_and_appendix(self):
        import io
        import docx

        content = StorylineContent(
            conversation_title="Q3 Funnel Review",
            turns=[StorylineTurn(answer_id="a1", question="Why did signups drop?",
                                  answer="Signups dropped 12% after the consent page.",
                                  facts=[], caveats=[], created_at="2026-08-15T00:00:00Z")],
            code_appendix=[CodeAppendixEntry(label="df_1", kind="sql", code="SELECT 1",
                                             source_answer_id="a1", is_dependency=False)],
            estimated_tokens=10, over_budget=False,
        )
        data = render_docx(content)
        self.assertIsInstance(data, bytes)
        doc = docx.Document(io.BytesIO(data))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Q3 Funnel Review", full_text)
        self.assertIn("Why did signups drop?", full_text)
        self.assertIn("Signups dropped 12%", full_text)
        self.assertIn("Code Appendix", full_text)
        self.assertIn("SELECT 1", full_text)

    def test_real_assembly_output_renders_its_dependency_sql_into_markdown(self):
        """The other renderer tests hand-build StorylineContent, so a mismatch
        between what assemble_storyline emits and what the renderers expect would be
        invisible. This chains the two for real."""
        conv = {"id": "c1", "title": "Q3 Funnel Review", "messages": [
            _msg("a1", question="SQL turn",
                 queries_run=["SELECT revenue FROM events WHERE region='EU'"],
                 produced_df_label="df_1"),
            _msg("a2", question="Python turn", answer="The total is 6",
                 python_cells=[{"code": "result = df_1['revenue'].sum()",
                                "df_label": "df_1", "result_summary": 6}]),
        ]}
        md = render_markdown(assemble_storyline(conv, ["a2"]))  # a1 NOT selected
        self.assertIn("# Q3 Funnel Review", md)
        self.assertIn("## Python turn", md)
        self.assertIn("SELECT revenue FROM events WHERE region='EU'", md)
        self.assertIn("included as a dependency of df_1", md)

    def test_note_entries_render_legibly_in_both_renderers(self):
        conv = {"id": "c1", "title": "Test", "messages": [
            _msg("a1", python_cells=[{"code": "result = df_1.sum()", "df_label": "df_1",
                                      "result_summary": 1}]),
        ]}
        content = assemble_storyline(conv, ["a1"])

        md = render_markdown(content)
        self.assertIn("(note)", md)
        self.assertIn("Provenance gap", md)
        self.assertIn("not recorded", md)

        import io
        import docx
        doc = docx.Document(io.BytesIO(render_docx(content)))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("(note)", full_text)
        self.assertIn("Provenance gap", full_text)
        self.assertIn("not recorded", full_text)

    def test_markdown_fence_is_longer_than_any_backtick_run_in_the_code(self):
        """M4. A ``` inside a docstring or SQL comment used to close the fence early
        and dump the rest of the document -- including later code -- out as prose."""
        content = StorylineContent(
            conversation_title="T",
            code_appendix=[CodeAppendixEntry(label="df_1", kind="python",
                                             code='x = """```"""',
                                             source_answer_id="a1", is_dependency=False)])
        md = render_markdown(content)
        self.assertIn("````python", md)
        self.assertIn('x = """```"""', md)
        # The fence still closes: exactly two four-backtick markers, nothing longer.
        self.assertEqual(md.count("````"), 2)

    def test_markdown_heading_collapses_a_multiline_question(self):
        content = StorylineContent(
            conversation_title="T",
            turns=[StorylineTurn(answer_id="a1", question="line one\nline two",
                                 answer="A", facts=[], caveats=[], created_at="")])
        self.assertIn("## line one line two", render_markdown(content))

    def test_created_at_is_rendered_in_both_formats(self):
        """M8. An audit artifact with no dates on any turn is a real gap."""
        content = StorylineContent(
            conversation_title="T",
            turns=[StorylineTurn(answer_id="a1", question="Q", answer="A", facts=[],
                                 caveats=[], created_at="2026-08-15T00:00:00Z")])
        self.assertIn("2026-08-15T00:00:00Z", render_markdown(content))

        import io
        import docx
        doc = docx.Document(io.BytesIO(render_docx(content)))
        self.assertIn("2026-08-15T00:00:00Z", "\n".join(p.text for p in doc.paragraphs))

    def test_render_docx_strips_control_characters_instead_of_raising(self):
        """M6. python-docx raises ValueError on any C0 control character, and
        LLM-authored answers / warehouse SQL can carry \\x0b or \\x0c."""
        content = StorylineContent(
            conversation_title="Title\x0bwith\x0ccontrols",
            turns=[StorylineTurn(answer_id="a1", question="bad \x0b heading",
                                 answer="ans \x0c body", facts=["fact \x01"],
                                 caveats=["caveat \x02"], created_at="2026-08-15T00:00:00Z")],
            code_appendix=[CodeAppendixEntry(label="df_1", kind="sql",
                                             code="SELECT \x0b 1", source_answer_id="a1",
                                             is_dependency=False)])
        import io
        import docx
        doc = docx.Document(io.BytesIO(render_docx(content)))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("\x0b", full_text)
        self.assertNotIn("\x0c", full_text)
        self.assertIn("SELECT  1", full_text)

    def test_render_docx_raises_typed_error_when_python_docx_is_missing(self):
        """I1. The import used to be at module scope, so a broken python-docx/lxml
        install made `import analytics_platform.api` fail and killed every endpoint.
        It is now lazy and raises a typed error the route maps to 503, leaving
        Markdown export and every other endpoint working."""
        content = StorylineContent(conversation_title="T")
        with patch.dict(sys.modules, {"docx": None}):
            with self.assertRaises(DocxRendererUnavailable):
                render_docx(content)
            self.assertIn("# T", render_markdown(content))
