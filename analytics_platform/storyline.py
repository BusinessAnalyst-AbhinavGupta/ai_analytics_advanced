"""Pure, DB-free assembly of a selective export ("storyline") from an already-fetched
stakeholder conversation dict. No I/O here -- Task 3/4's renderers and the API layer
own fetching and formatting; this module only decides WHAT goes into the export and
resolves the Code Appendix's cross-turn dependencies.
"""
import io
from bisect import bisect_left
from dataclasses import dataclass, field
from typing import TYPE_CHECKING, Any, Dict, List, Optional

if TYPE_CHECKING:            # narrative imports this module, so keep it one-way
    from .narrative import NarratedStoryline

CHARS_PER_TOKEN_ESTIMATE = 4
WARN_TOKEN_THRESHOLD = 50_000

# C0 control characters are illegal in XML 1.0 except tab/LF/CR, and python-docx
# raises ValueError("All strings must be XML compatible...") on any of them. LLM
# answers and warehouse-derived SQL can carry \x0b/\x0c, so every string handed to
# python-docx is scrubbed through _xml_safe first.
_CONTROL_TRANSLATION = {c: None for c in range(32) if c not in (9, 10, 13)}


class DocxRendererUnavailable(RuntimeError):
    """python-docx (or its native lxml dependency) could not be imported.

    Raised by render_docx only. Kept as a typed error so the API layer can map it
    to a 503 for the Word format alone instead of the import taking down the whole
    process at module load -- a missing optional renderer must not break Markdown
    export or any other endpoint.
    """


@dataclass
class StorylineTurn:
    answer_id: str
    question: str
    answer: str
    facts: List[str]
    caveats: List[str]
    created_at: str


@dataclass
class CodeAppendixEntry:
    label: str          # df_label this code relates to, "" for a plain SQL-only turn
    kind: str            # "sql" | "python" | "note"
    code: str
    source_answer_id: str
    is_dependency: bool  # True if pulled in only because a selected turn needs it


@dataclass
class StorylineContent:
    conversation_title: str
    turns: List[StorylineTurn] = field(default_factory=list)
    code_appendix: List[CodeAppendixEntry] = field(default_factory=list)
    estimated_tokens: int = 0
    over_budget: bool = False
    # How many distinct df_labels referenced by a selected Python turn could not be
    # traced back to a producing turn. Each one also gets a visible kind="note"
    # appendix entry; this count lets callers warn without re-scanning the appendix.
    unresolved_dependency_count: int = 0
    # Set only when an export asked to be narrated (Plan B Task 3). Renderers use
    # it when it is present and usable, and fall back to the turn-by-turn layout
    # otherwise -- narration is an enhancement, never a dependency.
    narrative: Optional["NarratedStoryline"] = None


def assemble_storyline(conversation: Dict[str, Any], answer_ids: List[str]) -> StorylineContent:
    selected_ids = set(answer_ids)
    all_messages = conversation.get("messages", [])

    # df_label is a *recyclable display label*, not an identity: the in-memory
    # ConversationDataCache is LRU-bounded and is lost on API restart, so the same
    # label can be issued twice within one persisted conversation. A last-wins
    # {label: message} dict therefore mis-attributes a Python turn to whichever
    # turn happened to reuse the label LAST -- possibly one that ran *after* it.
    # Resolution is nearest-preceding-producer instead: a Python cell can only
    # operate on a frame some earlier turn cached, so the label it saw is the most
    # recent prior producer of that label.
    label_positions: Dict[str, List[int]] = {}
    for i, m in enumerate(all_messages):
        label = m.get("produced_df_label")
        if label:
            label_positions.setdefault(label, []).append(i)

    def _producer_before(label: str, index: int) -> Optional[Dict[str, Any]]:
        positions = label_positions.get(label)
        if not positions:
            return None
        k = bisect_left(positions, index)
        if k == 0:
            return None
        return all_messages[positions[k - 1]]

    turns: List[StorylineTurn] = []
    appendix: List[CodeAppendixEntry] = []
    dependency_answer_ids_added: set = set()
    unresolved_labels: set = set()

    for index, m in enumerate(all_messages):
        if m["answer_id"] not in selected_ids:
            continue
        turns.append(StorylineTurn(
            answer_id=m["answer_id"], question=m["question"], answer=m["answer"],
            facts=list(m.get("facts", [])), caveats=list(m.get("caveats", [])),
            created_at=m.get("created_at", ""),
        ))
        for q in m.get("queries_run", []):
            appendix.append(CodeAppendixEntry(
                label=m.get("produced_df_label", ""), kind="sql", code=q,
                source_answer_id=m["answer_id"], is_dependency=False))
        for p in m.get("python_cells", []):
            label = p.get("df_label", "")
            appendix.append(CodeAppendixEntry(
                label=label, kind="python", code=p.get("code", ""),
                source_answer_id=m["answer_id"], is_dependency=False))
            if not label:
                continue
            dep_msg = _producer_before(label, index)
            if dep_msg is None:
                # Never drop this silently: every conversation created before
                # produced_df_label shipped backfills to "", so otherwise the whole
                # Code Appendix would quietly lose its provenance.
                if label not in unresolved_labels:
                    unresolved_labels.add(label)
                    appendix.append(CodeAppendixEntry(
                        label=label, kind="note",
                        code=(f"The source query for '{label}' is not recorded in this "
                              f"conversation, so its provenance could not be included."),
                        source_answer_id=m["answer_id"], is_dependency=True))
                continue
            if (dep_msg["answer_id"] not in selected_ids
                    and dep_msg["answer_id"] not in dependency_answer_ids_added):
                for q in dep_msg.get("queries_run", []):
                    appendix.append(CodeAppendixEntry(
                        label=label, kind="sql", code=q,
                        source_answer_id=dep_msg["answer_id"], is_dependency=True))
                dependency_answer_ids_added.add(dep_msg["answer_id"])

    estimate_text = "\n".join(
        t.question + t.answer + " ".join(t.facts) + " ".join(t.caveats) for t in turns
    ) + "\n".join(e.code for e in appendix)
    estimated_tokens = len(estimate_text) // CHARS_PER_TOKEN_ESTIMATE

    return StorylineContent(
        conversation_title=conversation.get("title", ""),
        turns=turns, code_appendix=appendix,
        estimated_tokens=estimated_tokens,
        over_budget=estimated_tokens > WARN_TOKEN_THRESHOLD,
        unresolved_dependency_count=len(unresolved_labels),
    )


def _one_line(text: str) -> str:
    """Collapse whitespace so a multi-line question can't turn a heading into prose."""
    return " ".join((text or "").split())


def _xml_safe(text: str) -> str:
    """Drop C0 control characters (except tab/LF/CR) that python-docx rejects."""
    return (text or "").translate(_CONTROL_TRANSLATION)


def _fence_for(code: str) -> str:
    """A fence at least one backtick longer than the longest run inside `code`, so a
    ``` in a docstring or SQL comment can't terminate the block early and spill the
    rest of the document out as prose."""
    longest = current = 0
    for ch in code or "":
        current = current + 1 if ch == "`" else 0
        longest = max(longest, current)
    return "`" * max(3, longest + 1)


def _narrated_markdown_head(content: StorylineContent) -> List[str]:
    """The narrated document's prose. Provenance is printed per section so a
    reader can jump from a claim to the turn that produced it."""
    n = content.narrative
    lines = [f"# {_one_line(n.title) or 'Storyline Export'}", ""]
    if n.executive_summary:
        lines += ["## Executive summary", "", n.executive_summary, ""]
    for section in n.sections:
        lines.append(f"## {_one_line(section.heading)}")
        lines.append("")
        lines.append(section.body)
        if section.answer_ids:
            lines += ["", f"*Based on: {', '.join(section.answer_ids)}*"]
        lines.append("")
    if n.caveats:
        lines += ["## Caveats", ""]
        lines += [f"- {c}" for c in n.caveats]
        lines.append("")
    return lines


def _turn_by_turn_markdown_head(content: StorylineContent) -> List[str]:
    lines = [f"# {_one_line(content.conversation_title) or 'Storyline Export'}", ""]
    for t in content.turns:
        lines.append(f"## {_one_line(t.question)}")
        lines.append("")
        if t.created_at:
            lines.append(f"*{t.created_at}*")
            lines.append("")
        lines.append(t.answer)
        if t.facts:
            lines.append("")
            lines.append("**Facts:** " + "; ".join(t.facts))
        if t.caveats:
            lines.append("")
            lines.append("**Caveats:** " + "; ".join(t.caveats))
        lines.append("")
    return lines


def render_markdown(content: StorylineContent) -> str:
    narrative = content.narrative
    lines = (_narrated_markdown_head(content) if narrative is not None and narrative.ok
             else _turn_by_turn_markdown_head(content))
    if content.code_appendix:
        lines.append("## Code Appendix")
        lines.append("")
        for e in content.code_appendix:
            heading = f"### {_one_line(e.label) or e.source_answer_id} ({e.kind})"
            if e.kind == "note":
                # No fence: a note is prose about a *missing* block, not code.
                lines.append(heading)
                lines.append("")
                lines.append(f"> **Provenance gap:** {_one_line(e.code)}")
                lines.append("")
                continue
            if e.is_dependency:
                heading += f" — (included as a dependency of {e.label})"
            fence = _fence_for(e.code)
            lines.append(heading)
            lines.append(f"{fence}{e.kind}")
            lines.append(e.code)
            lines.append(fence)
            lines.append("")
    return "\n".join(lines)


def _code_style(doc: Any) -> Any:
    """A real monospaced paragraph style, defined once, with the East-Asian font set
    too (setting run.font.name alone leaves w:eastAsia unset and Word falls back)."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn

    style = doc.styles.add_style("Storyline Code", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    return style


def render_docx(content: StorylineContent) -> bytes:
    # Imported lazily: python-docx pulls in native lxml, and a module-scope import
    # here would make an install/build failure take down every endpoint in api.py
    # rather than just the Word format. See DocxRendererUnavailable.
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - exercised via the typed error
        raise DocxRendererUnavailable(str(exc)) from exc

    doc = docx.Document()
    code_style = _code_style(doc)
    narrative = content.narrative
    if narrative is not None and narrative.ok:
        doc.add_heading(_xml_safe(_one_line(narrative.title)) or "Storyline Export", level=1)
        if narrative.executive_summary:
            doc.add_heading("Executive summary", level=2)
            doc.add_paragraph(_xml_safe(narrative.executive_summary))
        for section in narrative.sections:
            doc.add_heading(_xml_safe(_one_line(section.heading)), level=2)
            doc.add_paragraph(_xml_safe(section.body))
            if section.answer_ids:
                cite = doc.add_paragraph(
                    _xml_safe("Based on: " + ", ".join(section.answer_ids)))
                for run in cite.runs:
                    run.italic = True
        if narrative.caveats:
            doc.add_heading("Caveats", level=2)
            for caveat in narrative.caveats:
                doc.add_paragraph(_xml_safe(caveat), style="List Bullet")
    else:
        doc.add_heading(_xml_safe(_one_line(content.conversation_title)) or "Storyline Export", level=1)
        for t in content.turns:
            doc.add_heading(_xml_safe(_one_line(t.question)), level=2)
            if t.created_at:
                stamp = doc.add_paragraph(_xml_safe(t.created_at))
                for run in stamp.runs:
                    run.italic = True
            doc.add_paragraph(_xml_safe(t.answer))
            if t.facts:
                doc.add_paragraph(_xml_safe("Facts: " + "; ".join(t.facts)))
            if t.caveats:
                doc.add_paragraph(_xml_safe("Caveats: " + "; ".join(t.caveats)))
    if content.code_appendix:
        doc.add_heading("Code Appendix", level=1)
        for e in content.code_appendix:
            heading = f"{_one_line(e.label) or e.source_answer_id} ({e.kind})"
            if e.is_dependency and e.kind != "note":
                heading += f" — included as a dependency of {e.label}"
            doc.add_heading(_xml_safe(heading), level=3)
            if e.kind == "note":
                note_para = doc.add_paragraph(_xml_safe("Provenance gap: " + _one_line(e.code)))
                for run in note_para.runs:
                    run.italic = True
                continue
            doc.add_paragraph(_xml_safe(e.code), style=code_style)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
